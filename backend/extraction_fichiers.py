"""
extraction_fichiers.py — Extraction depuis PDF, Word, Excel, Images
Ralnejj Santé v3
"""
import os
import base64
import pypdf
from pathlib import Path


# ════════════════════════════════════════════════════════════════
# PDF
# ════════════════════════════════════════════════════════════════
def extraire_texte_pdf(chemin_fichier: str, max_chars: int = 8000) -> str:
    try:
        morceaux = []
        with open(chemin_fichier, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    morceaux.append(t.strip())

        texte_complet = "\n\n".join(morceaux).strip()
        if len(texte_complet) > max_chars:
            texte_complet = texte_complet[:max_chars] + "\n\n[...document tronqué...]"
        return texte_complet
    except Exception as e:
        print(f"[PDF ERREUR] {e}")
        return ""


# ════════════════════════════════════════════════════════════════
# WORD (.docx)
# ════════════════════════════════════════════════════════════════
def extraire_texte_word(chemin_fichier: str, max_chars: int = 8000) -> str:
    try:
        from docx import Document
        doc = Document(chemin_fichier)

        morceaux = []

        # Paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                morceaux.append(para.text.strip())

        # Tableaux
        for table in doc.tables:
            for row in table.rows:
                ligne = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if ligne:
                    morceaux.append(ligne)

        texte_complet = "\n\n".join(morceaux).strip()
        if len(texte_complet) > max_chars:
            texte_complet = texte_complet[:max_chars] + "\n\n[...document tronqué...]"
        return texte_complet
    except Exception as e:
        print(f"[WORD ERREUR] {e}")
        return ""


# ════════════════════════════════════════════════════════════════
# EXCEL (.xlsx / .xls)
# ════════════════════════════════════════════════════════════════
def extraire_texte_excel(chemin_fichier: str, max_chars: int = 8000) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(chemin_fichier, data_only=True)

        morceaux = []

        for nom_feuille in wb.sheetnames:
            ws = wb[nom_feuille]
            morceaux.append(f"=== Feuille : {nom_feuille} ===")

            for row in ws.iter_rows(values_only=True):
                # Filtrer les lignes complètement vides
                valeurs = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if valeurs:
                    morceaux.append(" | ".join(valeurs))

        texte_complet = "\n".join(morceaux).strip()
        if len(texte_complet) > max_chars:
            texte_complet = texte_complet[:max_chars] + "\n\n[...données tronquées...]"
        return texte_complet
    except Exception as e:
        print(f"[EXCEL ERREUR] {e}")
        return ""


# ════════════════════════════════════════════════════════════════
# IMAGE — encode en base64 pour l'API vision de Groq
# ════════════════════════════════════════════════════════════════
EXTENSIONS_IMAGES = {".jpg", ".jpeg", ".png", ".webp", ".gif"}

def encoder_image_base64(chemin_fichier: str) -> dict | None:
    """
    Encode une image en base64 pour l'envoyer à l'API vision.
    Retourne un dict {base64, media_type} ou None si erreur.
    """
    try:
        extension = Path(chemin_fichier).suffix.lower()
        media_types = {
            ".jpg":  "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png":  "image/png",
            ".webp": "image/webp",
            ".gif":  "image/gif",
        }
        media_type = media_types.get(extension, "image/jpeg")

        with open(chemin_fichier, "rb") as f:
            data = base64.b64encode(f.read()).decode("utf-8")

        return {"base64": data, "media_type": media_type}
    except Exception as e:
        print(f"[IMAGE ERREUR] {e}")
        return None


# ════════════════════════════════════════════════════════════════
# Fonction principale — détecte le type et extrait
# ════════════════════════════════════════════════════════════════
def extraire_fichier(chemin_fichier: str) -> dict:
    """
    Détecte le type de fichier et extrait le contenu.

    Retourne :
    {
        "type":       "texte" | "image",
        "contenu":    str (texte extrait) — si type == "texte"
        "image_data": dict {base64, media_type} — si type == "image"
        "erreur":     str | None
    }
    """
    extension = Path(chemin_fichier).suffix.lower()

    # Image → vision
    if extension in EXTENSIONS_IMAGES:
        image_data = encoder_image_base64(chemin_fichier)
        if image_data:
            return {"type": "image", "image_data": image_data, "erreur": None}
        return {"type": "image", "image_data": None, "erreur": "Impossible d'encoder l'image."}

    # PDF
    if extension == ".pdf":
        texte = extraire_texte_pdf(chemin_fichier)
        if texte:
            return {"type": "texte", "contenu": texte, "erreur": None}
        return {"type": "texte", "contenu": "", "erreur": "Impossible d'extraire le texte du PDF (scan/image ?)."}

    # Word
    if extension in {".docx", ".doc"}:
        texte = extraire_texte_word(chemin_fichier)
        if texte:
            return {"type": "texte", "contenu": texte, "erreur": None}
        return {"type": "texte", "contenu": "", "erreur": "Impossible d'extraire le texte du document Word."}

    # Excel
    if extension in {".xlsx", ".xls"}:
        texte = extraire_texte_excel(chemin_fichier)
        if texte:
            return {"type": "texte", "contenu": texte, "erreur": None}
        return {"type": "texte", "contenu": "", "erreur": "Impossible d'extraire les données du fichier Excel."}

    return {"type": "inconnu", "contenu": "", "erreur": f"Format non supporté : {extension}"}